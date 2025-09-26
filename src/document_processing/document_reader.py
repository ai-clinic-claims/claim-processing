import os
import logging
import tempfile
from typing import Dict, List, Any
import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract
import openpyxl
import pdfplumber

from utils.logger import setup_logger

logger = setup_logger(__name__)

class DocumentReader:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            '.doc': self._read_docx,
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.csv': self._read_csv,
            '.txt': self._read_text,
            '.jpg': self._read_image,
            '.jpeg': self._read_image,
            '.png': self._read_image,
            '.tiff': self._read_image
        }
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a file with metadata"""
        try:
            if not os.path.exists(file_path):
                return {'error': f'File not found: {file_path}'}
            
            file_ext = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path)
            
            if file_ext not in self.supported_formats:
                return {
                    'file_path': file_path,
                    'file_type': file_ext,
                    'file_size': file_size,
                    'content': f'Unsupported file format: {file_ext}',
                    'pages': 0,
                    'word_count': 0
                }
            
            # Extract content
            content = self.supported_formats[file_ext](file_path)
            
            # Calculate statistics
            word_count = len(content.split()) if content else 0
            pages = content.count('\f') + 1 if content else 1  # Estimate pages
            
            return {
                'file_path': file_path,
                'file_type': file_ext,
                'file_size': file_size,
                'content': content,
                'pages': pages,
                'word_count': word_count,
                'extraction_success': True
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'file_path': file_path,
                'file_type': file_ext,
                'file_size': file_size,
                'content': f'Error reading file: {str(e)}',
                'pages': 0,
                'word_count': 0,
                'extraction_success': False
            }
    
    def _read_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for best results"""
        text = ""
        
        try:
            # Method 1: Try pdfplumber first (better for complex PDFs)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed, trying PyPDF2: {str(e)}")
            
            # Method 2: Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {str(e)}")
            
            return text.strip() if text.strip() else "No text could be extracted from PDF"
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return f"Error extracting PDF content: {str(e)}"
    
    def _read_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading Word document {file_path}: {str(e)}")
            return f"Error extracting Word document content: {str(e)}"
    
    def _read_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            text = ""
            # Try reading with pandas first
            try:
                excel_file = pd.ExcelFile(file_path)
                for sheet_name in excel_file.sheet_names:
                    text += f"--- Sheet: {sheet_name} ---\n"
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    text += df.to_string() + "\n\n"
            except:
                # Fallback to openpyxl for complex files
                workbook = openpyxl.load_workbook(file_path)
                for sheet_name in workbook.sheetnames:
                    text += f"--- Sheet: {sheet_name} ---\n"
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows(values_only=True):
                        text += " | ".join(str(cell) if cell is not None else "" for cell in row) + "\n"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {str(e)}")
            return f"Error extracting Excel content: {str(e)}"
    
    def _read_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {str(e)}")
            return f"Error extracting CSV content: {str(e)}"
    
    def _read_text(self, file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return f"Error reading text file: {str(e)}"
    
    def _read_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            # Check if tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except:
                return "OCR not available. Install tesseract-ocr for image text extraction."
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading image file {file_path}: {str(e)}")
            return f"Error extracting text from image: {str(e)}"
    
    def extract_from_attachments(self, attachments: List[Dict]) -> Dict[str, Any]:
        """Extract text from all attachments"""
        results = {
            'total_attachments': len(attachments),
            'successful_extractions': 0,
            'failed_extractions': 0,
            'attachments': []
        }
        
        for attachment in attachments:
            file_path = attachment.get('path', '')
            if not file_path or not os.path.exists(file_path):
                results['attachments'].append({
                    'filename': attachment.get('filename', 'Unknown'),
                    'error': 'File path not found',
                    'extraction_success': False
                })
                results['failed_extractions'] += 1
                continue
            
            extraction_result = self.extract_text_from_file(file_path)
            extraction_result['filename'] = attachment.get('filename', 'Unknown')
            
            if extraction_result.get('extraction_success', False):
                results['successful_extractions'] += 1
            else:
                results['failed_extractions'] += 1
            
            results['attachments'].append(extraction_result)
        
        return results